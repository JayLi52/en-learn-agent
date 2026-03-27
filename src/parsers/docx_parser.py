"""DOCX 文档解析器"""
from typing import List
from pathlib import Path
from .base import BaseParser, Document


class DocxParser(BaseParser):
    """解析 DOCX 文档"""

    def parse(self) -> List[Document]:
        self.validate_file()
        documents = []

        from docx import Document as DocxDocument

        doc = DocxDocument(self.file_path)

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 合并为完整文档
        full_text = "\n\n".join(paragraphs)

        if full_text.strip():
            documents.append(Document(
                content=full_text,
                source=str(self.file_path),
                metadata={
                    "file_type": "docx",
                }
            ))

        return documents